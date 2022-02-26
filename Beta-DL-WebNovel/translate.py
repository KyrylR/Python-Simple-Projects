import asyncio
import urllib.parse
from difflib import SequenceMatcher
from time import time, sleep

import requests
from bs4 import BeautifulSoup
from docx import Document
from pyppeteer import launch


def performance(func):
    def wrapper(*args, **kwargs):
        t1 = time()
        result = func(*args, **kwargs)
        t2 = time()
        print(f'It took {t2 - t1} s')
        return result

    return wrapper


def start_scrap(url):
    chapter_text = []
    response = requests.get(url)
    soup = BeautifulSoup(response.text, 'html.parser')
    temp = soup.find_all('p')
    next_to_scrap = soup.select('.nav-next a')[0].get('href', None)
    for item in temp:
        # print("0" + str(item.parent.nextSibling) + "0")
        if SequenceMatcher(a=str(item.parent.nextSibling), b='.entry-content').ratio() > 0.6:
            if item.text == " ":
                continue
            chapter_text.append(item.text)

    return (chapter_text), next_to_scrap


def make_translated_data(chapter_list):
    text = ""
    res_list = []
    for count, item in enumerate(chapter_list.split("   ")):
        if item == '':
            continue
        text = text + item + '\n'
        if len(text) + len(item) + 1 > 1000:
            if count != len(chapter_list) - 1:
                res_list.append(text)
                text = item

    return res_list


translate_result = []


async def translate(to_translate):
    # Launch the browser
    browser = await launch()

    # Open a new browser page
    page = await browser.newPage()

    for item in to_translate:
        text_raw = item
        text_raw = urllib.parse.quote(text_raw, safe='~@#$&()*!+=:;,.?/\'', encoding=None, errors=None)
        url = f'https://www.deepl.com/translator#en/ru/{text_raw}'

        # Open web page by url
        await page.goto(url)
        text = ""
        while len(text) < 10:
            sleep(0.5)
            page_content = await page.content()
            # Process extracted content with BeautifulSoup
            soup = BeautifulSoup(page_content, features="lxml")
            text = soup.find(id="target-dummydiv").get_text()

        translate_result.append(text)
    # Close browser
    await browser.close()


def getText_fromDocx(filename):
    doc = Document(filename)
    fullText = []
    for para in doc.paragraphs:
        fullText.append(para.text)
    return '\n'.join(fullText)


def write_docx(str_list, title):
    document = Document()

    document.add_heading(f'{title}', 0)

    for item in str_list:
        document.add_paragraph(item)

    document.add_page_break()

    document.save(f'{title}_translated.docx')


@performance
def process(titles_list: list):
    for item in titles_list:
        global translate_result
        translate_result = []
        text = getText_fromDocx(f'data/{item}.docx')
        asyncio.get_event_loop().run_until_complete(translate(make_translated_data(text)))
        write_docx(translate_result, item)


if __name__ == "__main__":
    pass
