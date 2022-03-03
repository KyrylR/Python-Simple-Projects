from docx import Document


def writeToDocx(str_list, title, index=None):
    document = Document()

    title_name = title if index is None else f'#{index}-{title}'

    document.add_heading(f'{title_name}', 0)

    for item in str_list:
        document.add_paragraph(item)

    document.add_page_break()

    print(f'Save file: {title_name}')
    document.save(f'data_to_translate\{title_name}.docx')


def getTextFromDocx(filename):
    doc = Document(filename)
    fullText = []
    for para in doc.paragraphs:
        fullText.append(para.text)
    return '\n'.join(fullText)


if __name__ == "__main__":
    pass
