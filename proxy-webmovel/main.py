import asyncio
import json
from difflib import SequenceMatcher

from bs4 import BeautifulSoup
from pyppeteer import launch

from Proxy import Proxy

from docx import Document
from docx.shared import Inches


async def get_chapter_list(page):
    chapter_list = []
    chapter_list_titles = []
    await page.goto("https://www.webnovel.com/book/spirit-fox_22448045506616205/catalog")
    page_content = await page.content()
    # Process extracted content with BeautifulSoup
    soup = BeautifulSoup(page_content, features="lxml")
    print(soup.text)
    temp = soup.find_all('a')
    # next_to_scrap = soup.select('.nav-next a')[0].get('href', None)
    flag = False
    for item in temp:
        # print("0" + str(item.parent.nextSibling) + "0")
        if item.text == ' 1  Filthy Animal 5d  ' or flag:
            flag = True
            chapter_list_titles.append(item.attrs['title'])
            chapter_list.append('https://www.webnovel.com/' + str(item.attrs['href']))

    return chapter_list, chapter_list_titles


async def get_chapter_text(chapters: list, page, titles):
    for index, item in enumerate(chapters):
        text = []
        await page.goto(item)
        page_content = await page.content()
        # Process extracted content with BeautifulSoup
        soup = BeautifulSoup(page_content, features="lxml")
        # print(soup.text)
        temp = soup.find_all('p')
        for p in temp:
            if len(p.attrs) == 0:
                text_to_upload = p.text
                text_to_upload = text_to_upload.replace('\n', '').strip()
                if text_to_upload == 'Paragraph comments':
                    break
                text.append(text_to_upload)

        write_docx(text, titles[index], index)


def write_docx(str_list, title, title_num):
    document = Document()

    document.add_heading(f'{title}', 0)

    for item in str_list:
        document.add_paragraph(item)

    document.add_page_break()

    document.save(f'#{title_num + 1}-{title}.docx')


async def get_data_webnovel():
    # proxy = Proxy()
    # g = proxy.generator_function()
    while True:
        try:
            # dictionary = next(g)
            # next_proxy = next(iter(dictionary.values()))

            # Launch the browser
            # browser = await launch(
            #     {
            #         'ignoreHTTPSErrors': True,
            #         'args': [
            #             f'--proxy-server={next_proxy}',
            #             '--ignore-certificate-errors'
            #         ]
            #     }
            # )
            browser = await launch()

            # Open a new browser page
            page = await browser.newPage()

            # url_test = "https://icanhazip.com/"

            # Open web page by url
            # await page.goto(url_test)

            chapter_list, chapter_list_titles = await get_chapter_list(page)
            # with open('listfile.txt', 'w') as filehandle:
            #     json.dump(chapter_list, filehandle)

            with open('listfile.txt', 'r') as filehandle:
                chapter_list = json.load(filehandle)
            await get_chapter_text(chapter_list, page, chapter_list_titles)

            # Close browser
            await browser.close()
            break

        except StopIteration as err:
            print('StopIteration Error')
            break
        except Exception as e:
            print(e)
            await browser.close()


asyncio.get_event_loop().run_until_complete(get_data_webnovel())
